import os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\Francisco\Desktop\Ordenes de servicio para taller mecanico"
CAPS = r"C:\Users\Francisco\Desktop\Ordenes de servicio para taller mecanico\captura de pantalla"
OUT  = r"C:\Users\Francisco\Desktop\Ordenes de servicio para taller mecanico\TAREA4_SOLID_UIUX_Francisco_Diaz.docx"

# Ajustar rutas con caracteres especiales
BASE_REAL = r"C:\Users\Francisco\Desktop\Órdenes de servicio para taller mecánico"
CAPS_REAL = BASE_REAL + r"\captura de pantalla"
OUT_REAL  = BASE_REAL + r"\TAREA4_SOLID_UIUX_Francisco_Diaz.docx"

AZUL_OSC = RGBColor(0x1E,0x40,0xAF)
AZUL     = RGBColor(0x25,0x63,0xEB)
VERDE    = RGBColor(0x05,0x96,0x69)
ROJO     = RGBColor(0xDC,0x26,0x26)
AMBER    = RGBColor(0xD9,0x77,0x06)
MORADO   = RGBColor(0x7C,0x3A,0xED)
S900     = RGBColor(0x0F,0x17,0x2A)
S700     = RGBColor(0x33,0x41,0x55)
S500     = RGBColor(0x64,0x74,0x8B)
BLANCO   = RGBColor(0xFF,0xFF,0xFF)

doc = Document()
for section in doc.sections:
    section.top_margin=Cm(2); section.bottom_margin=Cm(2)
    section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)

def shd(cell, hex6):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hex6)
    tcPr.append(s)

def h1(t,c=None):
    p=doc.add_heading(t,1)
    if c:
        for r in p.runs: r.font.color.rgb=c
def h2(t,c=None):
    p=doc.add_heading(t,2)
    if c:
        for r in p.runs: r.font.color.rgb=c

def para(t='',bold=False,italic=False,color=None,sz=11,align=WD_ALIGN_PARAGRAPH.LEFT,sa=6):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    if t:
        r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
        if color: r.font.color.rgb=color
    return p

def mpara(parts,sz=11,align=WD_ALIGN_PARAGRAPH.LEFT,sa=6):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_after=Pt(sa)
    for t,b,i,c in parts:
        r=p.add_run(t); r.bold=b; r.italic=i; r.font.size=Pt(sz)
        if c: r.font.color.rgb=c
    return p

def bul(t,prefix=None,pc=None):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after=Pt(3)
    if prefix:
        r=p.add_run(prefix); r.bold=True; r.font.size=Pt(11)
        if pc: r.font.color.rgb=pc
    r2=p.add_run(t); r2.font.size=Pt(11)

def callout(t,bg='ECFDF5'):
    tb=doc.add_table(rows=1,cols=1); tb.style='Table Grid'
    c=tb.cell(0,0); shd(c,bg)
    r=c.paragraphs[0].add_run(t); r.font.size=Pt(10)
    doc.add_paragraph()

def img_labeled(fname,label,bg='EFF6FF',fc='2563EB',w=5.0,cap=None):
    path=os.path.join(CAPS_REAL,fname)
    tb=doc.add_table(rows=2,cols=1); tb.style='Table Grid'
    c0=tb.cell(0,0); shd(c0,bg)
    r=c0.paragraphs[0].add_run(label); r.bold=True; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor.from_string(fc)
    c1=tb.cell(1,0); p=c1.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path,width=Inches(w))
    else:
        p.add_run('[Imagen: '+fname+']')
    if cap:
        pc=doc.add_paragraph(cap); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r2 in pc.runs:
            r2.font.size=Pt(9); r2.italic=True; r2.font.color.rgb=S500
    doc.add_paragraph()

def img_ba(fb,fa,cb='ANTES',ca='DESPUES'):
    path_b=os.path.join(CAPS_REAL,fb); path_a=os.path.join(CAPS_REAL,fa)
    tb=doc.add_table(rows=2,cols=2); tb.style='Table Grid'
    for idx,(lbl,bg,fc) in enumerate([('X  '+cb,'FEF2F2','DC2626'),('OK '+ca,'ECFDF5','059669')]):
        c=tb.cell(0,idx); shd(c,bg)
        r=c.paragraphs[0].add_run(lbl); r.bold=True; r.font.size=Pt(9)
        r.font.color.rgb=RGBColor.from_string(fc)
    for idx,path in enumerate([path_b,path_a]):
        c=tb.cell(1,idx); p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(path):
            p.add_run().add_picture(path,width=Inches(2.7))
        else:
            p.add_run('['+os.path.basename(path)+']')
    doc.add_paragraph()

def code_ba(before,after,tb_='ANTES',ta_='DESPUES'):
    tb=doc.add_table(rows=2,cols=2); tb.style='Table Grid'
    for idx,(tt,bg,fc) in enumerate([(tb_,'FEF2F2','DC2626'),(ta_,'ECFDF5','059669')]):
        c=tb.cell(0,idx); shd(c,bg)
        r=c.paragraphs[0].add_run(tt); r.bold=True; r.font.size=Pt(9)
        r.font.color.rgb=RGBColor.from_string(fc)
    for idx,code in enumerate([before,after]):
        c=tb.cell(1,idx); shd(c,'1E293B')
        r=c.paragraphs[0].add_run(code); r.font.name='Consolas'
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(0xE2,0xE8,0xF0)
    doc.add_paragraph()

def trow(t,cols,hdr_bg='1E293B',hdr_cols=None):
    tb=doc.add_table(rows=1,cols=len(cols)); tb.style='Table Grid'
    for i,h in enumerate(cols):
        c=tb.cell(0,i); shd(c,hdr_bg)
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10)
        r.font.color.rgb=BLANCO
    for row_data in t:
        row=tb.add_row()
        for i,val in enumerate(row_data):
            row.cells[i].paragraphs[0].add_run(str(val)).font.size=Pt(10)
    doc.add_paragraph()
    return tb

print("Iniciando generacion del documento Word...")

# ============================================================
# PORTADA
# ============================================================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after=Pt(4)
r=p.add_run('TAREA 4 — Principios SOLID y Mejoras UI/UX')
r.bold=True; r.font.size=Pt(20); r.font.color.rgb=AZUL_OSC

p2=doc.add_paragraph('Asignacion Practica — Ingenieria de Software II (INF-259)')
p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r2 in p2.runs: r2.font.size=Pt(13); r2.font.color.rgb=S700

mpara([('Estudiante: ',True,False,S900),('Francisco R. Diaz — Matricula: 2023-3300016',False,False,S700)],
      align=WD_ALIGN_PARAGRAPH.CENTER,sa=3)
mpara([('Universidad: ',True,False,S900),('UNEV — Universidad Nacional Evangelica',False,False,S700)],
      align=WD_ALIGN_PARAGRAPH.CENTER,sa=3)
mpara([('Proyecto: ',True,False,S900),('SIGEST — Sistema de Ordenes de Servicio para Taller Mecanico',False,False,S700)],
      align=WD_ALIGN_PARAGRAPH.CENTER,sa=3)
mpara([('Stack: ',True,False,S900),('Node.js + Express + Prisma + PostgreSQL | React + Vite + Tailwind CSS',False,False,S700)],
      align=WD_ALIGN_PARAGRAPH.CENTER,sa=3)
mpara([('Docente: ',True,False,S900),('Ing. Carlos Artemio Escalante Lorenzo',False,False,S700),('    Julio 2026',False,False,S500)],
      align=WD_ALIGN_PARAGRAPH.CENTER,sa=16)

doc.add_page_break()
print("Portada OK")

# ============================================================
# TEMA 1 — SOLID
# ============================================================
h1('Tema 1 — Principios SOLID en la Logica de Negocio',AZUL_OSC)
para('El sistema SIGEST implementa Arquitectura Hexagonal con DDD. Se muestran los 5 principios con fragmentos de codigo ANTES/DESPUES y los olores de codigo eliminados.',sa=10)

# SRP
h2('SRP — Single Responsibility Principle (Responsabilidad Unica)',AZUL)
mpara([('Regla: ',True,False,S900),('Cada clase debe tener una sola razon para cambiar.',False,False,S700)])
mpara([('Problema: ',True,False,ROJO),('El controlador vehiculo.routes.js combinaba validacion, transformacion y persistencia (3 razones para cambiar).',False,False,S700)])
code_ba(
'// router.post("/") mezclaba todo:\nconst {placa,marca,modelo,anio}=req.body;\n// VALIDAR en controlador\nif(!placa||!marca)\n  return res.status(400).json({error});\n// TRANSFORMAR datos aqui\nconst anioInt=parseInt(anio);\n// PERSISTIR con Prisma directo\nconst v=await prisma.vehiculo.create({...});',
'// CrearVehiculoUseCase — 1 responsabilidad\nclass CrearVehiculoUseCase {\n  constructor(vehiculoRepo){\n    this.repo=vehiculoRepo;\n  }\n  async ejecutar(datos){\n    if(!datos.placa||!datos.marca)\n      throw new Error("Obligatorios");\n    const anio=parseInt(datos.anio);\n    return this.repo.crear({...datos,anio});\n  }\n}\n// Controlador: solo delega\nrouter.post("/",(req,res)=>{\n  facade.crearVehiculo(req.body)\n    .then(v=>res.json(v));\n});',
'ANTES — vehiculo.routes.js (God Method)',
'DESPUES — CrearVehiculoUseCase.js'
)
callout('Olor eliminado: God Method — controlador HTTP mezclaba 3 responsabilidades. Ahora separadas en UseCase (logica) + Repository (persistencia) + Controlador (HTTP).')

# OCP
h2('OCP — Open/Closed Principle (Abierto/Cerrado)',VERDE)
mpara([('Regla: ',True,False,S900),('Abierto a la extension, cerrado a la modificacion.',False,False,S700)])
mpara([('Problema: ',True,False,ROJO),('Calculos de descuento con if/else — agregar nuevo descuento requeria modificar clase ya probada.',False,False,S700)])
code_ba(
'// Si/else en el calculador (violacion OCP)\nfunction calcularDesc(sub,tipo){\n  if(tipo==="frecuente")\n    return sub*0.10;\n  else if(tipo==="corporativo")\n    return sub*0.15;\n  else if(tipo==="promocion")\n    return sub*0.05;\n  // Para nuevo descuento: modificar aqui\n  return 0;\n}',
'// Patron Strategy — OCP aplicado\nclass CalculadoraOrden {\n  constructor(estrategia=new SinDescuento()){\n    this._e=estrategia;\n  }\n  calcularTotal(svc,rep){\n    const sub=/* suma */;\n    return { sub, desc:this._e.calcular(sub) };\n  }\n}\n// Nuevo descuento = clase nueva (no modifica):\nclass DescuentoCorporativo {\n  calcular(s){ return s*0.15; }\n  descripcion(){ return "Corporativo 15%"; }\n}',
'ANTES — if/else en calculador',
'DESPUES — Patron Strategy (OCP)'
)
callout('Olor eliminado: Cascada if/else — ahora se agregan nuevos descuentos creando clases nuevas. CalculadoraOrden nunca se modifica.')

# LSP
h2('LSP — Liskov Substitution Principle (Sustitucion de Liskov)',MORADO)
mpara([('Regla: ',True,False,S900),('Una subclase puede reemplazar a su clase base sin alterar el comportamiento.',False,False,S700)])
mpara([('Problema: ',True,False,ROJO),('Sin patron State, las transiciones eran if/else dispersos sobre el campo estado.',False,False,S700)])
code_ba(
'// Transiciones con if/else (sin LSP)\nfunction avanzar(orden){\n  if(orden.estado==="RECIBIDA")\n    orden.estado="EN_DIAGNOSTICO";\n  else if(orden.estado==="EN_DIAGNOSTICO")\n    orden.estado="PRESUPUESTADA";\n  // Cada nuevo estado: otro else-if\n  // Ninguno es sustituible\n}',
'// Patron State + LSP\nclass EstadoOrden {\n  avanzar(o){ throw Error("No impl"); }\n  rechazar(o){ throw Error("No impl"); }\n  nombre()  { throw Error("No impl"); }\n}\nclass EstadoRecibida extends EstadoOrden {\n  avanzar(o){ o._cambiarEstado(new EstadoEnDiagnostico()); }\n  nombre(){ return "RECIBIDA"; }\n}\nclass EstadoPresupuestada extends EstadoOrden {\n  avanzar(o) { o._cambiarEstado(new EstadoAprobada()); }\n  rechazar(o){ o._cambiarEstado(new EstadoRechazada()); }\n  nombre(){ return "PRESUPUESTADA"; }\n}',
'ANTES — if/else sin sustitucion',
'DESPUES — 10 clases Estado intercambiables'
)
callout('Olor eliminado: Switch/if-else sobre tipo — 10 clases de estado (EstadoRecibida, EstadoEnDiagnostico, EstadoPresupuestada, etc.) implementan la misma interfaz EstadoOrden y son intercambiables.')

# ISP
h2('ISP — Interface Segregation Principle (Segregacion de Interfaces)',AMBER)
mpara([('Regla: ',True,False,S900),('Preferir interfaces pequenas y especificas.',False,False,S700)])
mpara([('Problema: ',True,False,ROJO),('Si los Use Cases recibieran el TallerFacade completo (20+ metodos), tendrian acceso a metodos que no usan.',False,False,S700)])
code_ba(
'// UseCase recibe TODO el facade\nclass CrearOrdenUseCase {\n  constructor(facade){\n    this.facade=facade;\n    // Tiene acceso a facade.login()\n    // Tiene acceso a facade.eliminarOrden()\n    // Tiene acceso a facade.gestionarUsuario()\n    // Solo necesita: ordenRepo, clienteRepo\n  }\n}',
'// Solo las dependencias minimas necesarias\nclass CrearOrdenUseCase {\n  constructor(ordenRepo,clienteRepo,\n              vehiculoRepo,eventBus){\n    this._or=ordenRepo;\n    this._cr=clienteRepo;\n    this._vr=vehiculoRepo;\n    this._eb=eventBus;\n    // No conoce facturas, usuarios, login\n  }\n}\n// GenerarFacturaUseCase — interfaz diferente\nclass GenerarFacturaUseCase {\n  constructor(ordenRepo,facturaRepo,eb){...}\n}',
'ANTES — Depende del Facade completo',
'DESPUES — Solo recibe lo que necesita'
)
callout('Olor eliminado: Fat Interface — cada UseCase recibe exactamente los repositorios que necesita. container.js inyecta manualmente las 4 dependencias de cada caso de uso.')

# DIP
h2('DIP — Dependency Inversion Principle (Inversion de Dependencias)',ROJO)
mpara([('Regla: ',True,False,S900),('Depender de abstracciones, no de implementaciones concretas.',False,False,S700)])
mpara([('Problema: ',True,False,ROJO),('Controladores instanciaban PrismaClient directamente — acoplamiento total a la BD.',False,False,S700)])
code_ba(
'// Acoplamiento directo a Prisma\nconst {PrismaClient}=require("@prisma/client");\nconst prisma=new PrismaClient();\n\nrouter.post("/",(req,res)=>{\n  // Imposible hacer unit tests\n  // Imposible cambiar a MongoDB\n  const o=await prisma.orden.create(\n    {data:req.body}\n  );\n  res.json(o);\n});',
'// container.js — Inyeccion de dependencias\nconst ordenRepo=\n  new PrismaOrdenRepository();\nconst clienteRepo=\n  new PrismaClienteRepository();\nconst eventBus=new EventBus();\n\nconst tallerFacade=new TallerFacade({\n  ordenRepo,    // abstraccion\n  clienteRepo,  // abstraccion\n  vehiculoRepo, // abstraccion\n  eventBus,     // abstraccion\n});\n// Para MongoDB: solo nuevo repositorio\n// Para tests: inyectar InMemoryRepo',
'ANTES — new PrismaClient() en controlador',
'DESPUES — container.js + abstracciones'
)
callout('Olor eliminado: Hard-coded dependency — el container.js wirea las dependencias. Las capas superiores dependen de OrdenRepository (abstracto), no de PrismaOrdenRepository (concreto).')

doc.add_page_break()
print("SOLID OK")

# ============================================================
# TEMA 2 — UI/UX
# ============================================================
h1('Tema 2 — Mejoras Profesionales de UI/UX',AZUL_OSC)
para('Se implementaron 8 tecnicas de interfaz de usuario (minimo requerido: 6), cada una con justificacion tecnica y capturas reales del sistema SIGEST desplegado en produccion.',sa=10)

# T1: Wizard
h2('Tecnica 1 — Asistente / Wizard por Pasos — NuevaOrden.jsx',AZUL)
mpara([('Problema: ',True,False,ROJO),('Formulario unico saturado con todos los campos a la vez (cliente, vehiculo, prioridad, mecanico, servicios, notas) — sobrecarga cognitiva.',False,False,S700)])
bul('Wizard de 3 pasos: Paso 1 (Cliente y Vehiculo), Paso 2 (Servicios), Paso 3 (Resumen).')
bul('Barra de progreso visual con estados: completado, activo y pendiente.')
bul('Validacion por etapa — no se avanza sin completar el paso actual.')
bul('Subtotal en tiempo real visible en el Paso 2.')
mpara([('Principio: ',True,False,AZUL),('Progressive Disclosure — reduccion de carga cognitiva y prevencion de errores.',False,False,S700)])

para('ANTES — Formulario unico saturado:',bold=True,color=ROJO,sa=4)
img_labeled('Captura de pantalla 2026-07-10 005222.png',
    'ANTES: Formulario Nueva Orden — todos los campos en una sola pantalla sin guia',
    bg='FEF2F2',fc='DC2626',w=5.0,
    cap='Todos los campos visibles a la vez: Cliente, Vehiculo, Prioridad, Mecanico, Servicios, Notas.')

para('DESPUES — Wizard Paso 1: Cliente y Vehiculo:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-07-10 003124.png',
    'DESPUES Paso 1: Seleccion de Cliente, Vehiculo y Mecanico',
    bg='ECFDF5',fc='059669',w=4.5,
    cap='Paso 1: Ana Perez (001-0000002-2) + Honda Civic B789012 + Mecanico: Carlos Rodriguez.')

para('DESPUES — Wizard Paso 2: Servicios con subtotal:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-07-10 010205.png',
    'DESPUES Paso 2: Servicios seleccionados con subtotal RD$4,000.00 en tiempo real',
    bg='ECFDF5',fc='059669',w=4.5,
    cap='Barra de progreso visible (Paso 1 completado, Paso 2 activo). Servicio "de rutina" RD$4,000.00.')

para('DESPUES — Wizard Paso 3: Resumen y Confirmacion:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-07-10 010253.png',
    'DESPUES Paso 3: Resumen completo antes de confirmar la creacion de la orden',
    bg='ECFDF5',fc='059669',w=4.5,
    cap='Resumen: cliente, vehiculo, mecanico, servicios y total RD$4,000.00 antes del boton Crear Orden.')

# T2: Modal
h2('Tecnica 2 — Ventanas Modales — ConfirmModal.jsx',AZUL)
mpara([('Problema: ',True,False,ROJO),('window.alert() y window.confirm() bloquean el hilo JS y tienen apariencia inconsistente entre navegadores.',False,False,S700)])
bul('ConfirmModal.jsx — dialogo estilizado para confirmar eliminar y generar factura.')
bul('Overlay oscuro que centra la atencion sin perder el contexto de la lista.')
bul('Modal de edicion de cliente con campos validados y seccion colapsable para agregar vehiculo.')
bul('Botones con semantica: rojo = destruir, azul/verde = confirmar.')
mpara([('Principio: ',True,False,AZUL),('Prevencion de errores — confirmacion explicita antes de acciones irreversibles (Nielsen #5).',False,False,S700)])

para('Evidencia — Modal Editar Cliente:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-06-10 160326.png',
    'Modal Editar Informacion del Cliente — campos validados, seccion colapsable Agregar Vehiculo',
    bg='ECFDF5',fc='059669',w=3.5,
    cap='Modal con Nombre, Cedula/RNC, Telefono, Correo, Direccion y seccion colapsable Agregar Vehiculo (Opcional).')

# T3: Color
h2('Tecnica 3 — Paleta de Color Profesional — tailwind.config.js',AZUL)
mpara([('Problema: ',True,False,ROJO),('Sin paleta definida, los colores eran inconsistentes.',False,False,S700)])
trow([
    ('Primario Brand (sky-500)','#0EA5E9','Identidad visual y confianza'),
    ('Accion / Avanzar (blue-600)','#2563EB','Accion principal del sistema'),
    ('Exito / Guardar (emerald-600)','#059669','Confirmacion y exito'),
    ('Error / Eliminar (red-600)','#DC2626','Peligro y urgencia'),
    ('Advertencia (amber-600)','#D97706','Precaucion sin alarma'),
    ('Sidebar (slate-900)','#0F172A','Fondo del menu lateral'),
    ('Fondo general (slate-50)','#F8FAFC','Fondo neutro de contenido'),
],['Color / Uso','HEX','Justificacion'])
callout('Contraste WCAG AA verificado. Definido como brand colors en tailwind.config.js y usado consistentemente en todos los componentes.','EFF6FF')

# T4: DataGrid
h2('Tecnica 4 — DataGrid Profesional — Ordenes.jsx / Facturacion.jsx',AZUL)
mpara([('Problema: ',True,False,ROJO),('Lista simple de cards sin encabezados, sin paginacion, sin ordenamiento.',False,False,S700)])
bul('Encabezados de columna (Cliente/Vehiculo, Estado, Fecha) con ordenamiento al hacer clic.')
bul('Zebra striping — filas alternadas para mejor legibilidad horizontal.')
bul('Paginacion local — 10 registros por pagina con navegacion numerica.')
bul('Formato condicional — badges de color por estado (FACTURADA=verde, URGENTE=rojo).')
bul('DataGrid de Facturacion con columnas: Factura/Orden, Subtotal, ITBIS 18%, Total.')
mpara([('Principio: ',True,False,AZUL),('Consistencia y estandares — Heuristica de Nielsen #4.',False,False,S700)])

para('ANTES vs DESPUES — Lista de Ordenes:',bold=True,color=ROJO,sa=4)
img_ba('Captura de pantalla 2026-07-02 234255.png','Captura de pantalla 2026-07-10 003933.png',
       'ANTES: Lista cards simple sin encabezados',
       'DESPUES: DataGrid con encabezados, sort y paginacion')

para('DataGrid de Facturacion con ITBIS 18%:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-06-10 163337.png',
    'Panel Facturacion & Caja — DataGrid: Factura/Orden · Cliente/Vehiculo · Subtotal · ITBIS 18% · Total',
    bg='ECFDF5',fc='059669',w=5.5,
    cap='FAC-2D26929C: Subtotal RD$5,000 + ITBIS 18% RD$900 = Total RD$5,900. Buscador y boton Imprimir Reporte.')

# T5: Roles
h2('Tecnica 5 — Ocultacion de Opciones por Rol — navigationConfig.js',AZUL)
mpara([('Problema: ',True,False,ROJO),('Mostrar modulos inaccesibles genera confusion y riesgo de seguridad.',False,False,S700)])
bul('getNavItemsByRol(rol) filtra el sidebar segun el rol del usuario autenticado.')
bul('ADMIN: 6 modulos (Dashboard, Ordenes, Clientes, Usuarios, Equipo Tecnico, Facturacion).')
bul('MECANICO: solo Dashboard y Ordenes.')
bul('CAJERO: Dashboard, Ordenes, Clientes y Facturacion.')
bul('Boton "Eliminar Orden" visible solo para ADMIN.')
trow([
    ('Dashboard','✓','✓','✓','✓'),
    ('Ordenes','✓','✓','✓','✓'),
    ('Clientes','✓','✓','✗','✓'),
    ('Usuarios','✓','✗','✗','✗'),
    ('Equipo Tecnico','✓','✗','✗','✗'),
    ('Facturacion','✓','✓','✗','✓'),
    ('Eliminar Orden','✓','✗','✗','✗'),
],['Modulo','ADMIN','RECEPCIONISTA','MECANICO','CAJERO'])

para('Evidencia — Dashboard con sidebar completo (rol ADMIN):',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-06-10 162707.png',
    'Vista ADMIN: 6 modulos en sidebar + badge amarillo ADMIN + indicador Sistema en linea',
    bg='ECFDF5',fc='059669',w=5.5,
    cap='Francisco Diaz (Admin): ve Dashboard, Ordenes, Clientes, Usuarios, Equipo Tecnico, Facturacion. Badge ADMIN visible.')

# T6: Toast / Feedback
h2('Tecnica 6 — Retroalimentacion Visual — Toast.jsx + Badges',AZUL)
mpara([('Problema: ',True,False,ROJO),('window.alert() bloquea JS, apariencia inconsistente entre navegadores.',False,False,S700)])
bul('Toast.jsx — notificaciones flotantes animadas (slide-in) que desaparecen en 3 segundos.')
bul('Tres variantes: success (verde), error (rojo), info (azul).')
bul('Spinners animate-spin en botones de submit durante carga.')
bul('Badges de estado por color en toda la interfaz: FACTURADA, URGENTE, RECIBIDA.')
bul('Indicador "Sistema en linea" en el dashboard.')
mpara([('Principio: ',True,False,AZUL),('Visibilidad del estado del sistema — Heuristica de Nielsen #1.',False,False,S700)])

para('Evidencia — Badges de estado y DataGrid:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-07-10 003933.png',
    'Badges FACTURADA (verde #059669) y URGENTE (rojo #DC2626) con icono de alerta',
    bg='ECFDF5',fc='059669',w=5.5,
    cap='DataGrid de ordenes: badge FACTURADA verde, URGENTE rojo con alerta, fecha formateada en columna separada.')

# T7: Sidebar colapsable
h2('Tecnica 7 — Menu Lateral Colapsable — Layout.jsx',AZUL)
mpara([('Problema: ',True,False,ROJO),('Sidebar fijo de 256px reducia el espacio de contenido en pantallas medianas.',False,False,S700)])
bul('Boton toggle para colapsar/expandir el sidebar.')
bul('Modo colapsado: 72px con solo iconos (sin texto).')
bul('Tooltips al pasar el mouse sobre iconos colapsados (accesibilidad).')
bul('Estado persistido en localStorage.')
bul('Transicion suave: transition-all duration-300.')
mpara([('Principio: ',True,False,AZUL),('Flexibilidad y eficiencia — Heuristica de Nielsen #7.',False,False,S700)])

para('Evidencia — Sidebar expandido y colapsado:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-06-10 162543.png',
    'Sidebar expandido: iconos + texto, item Clientes activo (verde esmeralda), logo SIGEST arriba',
    bg='ECFDF5',fc='059669',w=5.5,
    cap='Sidebar expandido con Dashboard, Ordenes, Clientes (activo), Usuarios, Equipo Tecnico, Facturacion.')

# T8: Header dinamico
h2('Tecnica 8 — Header con Titulo Dinamico — Layout.jsx',AZUL)
mpara([('Problema: ',True,False,ROJO),('Sin indicacion clara de la seccion activa, el usuario podia perderse en el sistema.',False,False,S700)])
bul('Header superior fijo con backdrop-blur semitransparente.')
bul('Titulo dinamico que cambia segun la pagina actual (breadcrumb).')
bul('Boton hamburguesa en movil para abrir el sidebar.')
bul('Titulo con gradiente de color de marca.')
mpara([('Principio: ',True,False,AZUL),('Reconocimiento antes que recuerdo — Heuristica de Nielsen #6.',False,False,S700)])

para('Evidencia — Modulos con breadcrumb en header:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-06-10 163248.png',
    'Modulo Usuarios: breadcrumb "Usuarios" en header + tarjetas con roles ADMIN/RECEPCIONISTA',
    bg='ECFDF5',fc='059669',w=5.5,
    cap='Gestion de Personal: 3 usuarios (2 Admin, 1 Recepcionista). Header muestra "Usuarios" como titulo dinamico.')

para('Evidencia — Modulo Equipo Tecnico:',bold=True,color=VERDE,sa=4)
img_labeled('Captura de pantalla 2026-06-10 163320.png',
    'Modulo Equipo de Mecanicos: breadcrumb "Equipo Tecnico" + tarjetas con especialidad',
    bg='ECFDF5',fc='059669',w=5.5,
    cap='Equipo de Mecanicos: Carlos Rodriguez (Motor y Transmision), juan (Frenos). Header muestra "Equipo Tecnico".')

doc.add_page_break()
print("UI/UX OK")

# ============================================================
# CHECKLIST
# ============================================================
h1('Lista de Verificacion de Tecnicas Aplicadas',AZUL_OSC)
trow([
    ('1','Asistente / Wizard por pasos','NuevaOrden.jsx — 3 pasos con barra de progreso'),
    ('2','Ventanas modales (Modal / Dialog)','ConfirmModal.jsx — confirmar eliminar y facturar'),
    ('3','Tecnica de color profesional','tailwind.config.js — paleta brand + 7 colores semanticos'),
    ('4','DataGrid profesional','Ordenes.jsx — zebra, sort, paginacion, badges'),
    ('5','Ocultacion por rol/permiso','navigationConfig.js — filtrado por ADMIN/MECANICO/CAJERO'),
    ('6','Retroalimentacion visual (Toast, badges)','Toast.jsx — notificaciones + spinners + indicadores'),
    ('7','Menu lateral colapsable','Layout.jsx — toggle, tooltips, localStorage'),
    ('8','Header con titulo dinamico','Layout.jsx — breadcrumb dinamico, hamburguesa movil'),
],['#','Tecnica','Evidencia (pantalla / modulo)'])

callout('TOTAL: 8 tecnicas implementadas (minimo requerido: 6). Cada una justificada con las Heuristicas de Nielsen y evidencia fotografica del sistema real en produccion.','EFF6FF')
doc.add_page_break()

# ============================================================
# FUNCIONAMIENTO DEL PROYECTO
# ============================================================
h1('Funcionamiento del Proyecto',AZUL_OSC)
mpara([('URL produccion: ',True,False,S900),('https://taller-mecanico-frontend-khc6.onrender.com',False,False,AZUL)])
mpara([('GitHub: ',True,False,S900),('https://github.com/francisco-inger/taller-mecanico',False,False,AZUL)])
doc.add_paragraph()

h2('Stack Tecnologico',S700)
trow([
    ('Backend','Node.js + Express','API REST'),
    ('ORM','Prisma','Mapeo a PostgreSQL'),
    ('Base de datos','PostgreSQL 16','Persistencia relacional'),
    ('Frontend','React + Vite','Interfaz web SPA'),
    ('Estilos','Tailwind CSS','Utilidades CSS'),
    ('Autenticacion','JWT (jsonwebtoken)','Tokens seguros'),
    ('Contenedores','Docker + docker-compose','Despliegue reproducible'),
    ('Hosting','Render.com','Entorno de produccion'),
],['Capa','Tecnologia','Proposito'],'334155')

h2('Patrones de Diseno Implementados',S700)
trow([
    ('State','EstadoOrden (10 estados)','Ciclo de vida de la orden'),
    ('Strategy','CalculadoraOrden + 4 estrategias','Descuentos intercambiables'),
    ('Builder','OrdenServicioBuilder','Construccion de ordenes complejas'),
    ('Factory Method','ServicioFactory','Creacion de tipos de servicio'),
    ('Observer','EventBus','Notificaciones y auditoria'),
    ('Facade','TallerFacade','API simplificada del modulo'),
    ('Repository','OrdenRepository (abstracto)','Abstraccion de persistencia'),
    ('Decorator','ServicioDecorator','Extension de comportamiento'),
],['Patron','Clase','Proposito'],'334155')

h2('Evidencia — Sistema en Produccion',VERDE)

para('Login — Pantalla de inicio de sesion:',bold=True,color=AZUL,sa=4)
img_labeled('Captura de pantalla 2026-07-02 233905.png',
    'Login SIGEST — taller-mecanico-frontend-khc6.onrender.com',
    bg='EFF6FF',fc='2563EB',w=5.5,
    cap='Imagen de fondo tematica (motor), formulario glassmorphism, gradiente de marca verde-azul.')

para('Dashboard — Panel principal:',bold=True,color=AZUL,sa=4)
img_labeled('Captura de pantalla 2026-06-10 162707.png',
    'Dashboard: KPIs (Ordenes Activas, En Espera, Urgentes, Total), Actividad Reciente, Equipo Tecnico',
    bg='EFF6FF',fc='2563EB',w=5.5,
    cap='Sistema en linea. 0 activas, 1 urgente, 2 total. Tabla de actividad reciente y panel de carga del equipo.')

para('Detalle de Orden de Servicio:',bold=True,color=AZUL,sa=4)
img_labeled('Captura de pantalla 2026-06-10 162838.png',
    'Ordenes de Servicio con panel lateral — detalle completo: cliente, vehiculo, mecanico, servicios',
    bg='EFF6FF',fc='2563EB',w=5.5,
    cap='Orden Juan Martinez / Toyota Corolla — FACTURADA URGENTE — Reparacion tren delantero RD$5,000.00.')

doc.add_page_break()

# ============================================================
# RUBRICA
# ============================================================
h1('Rubrica de Autoevaluacion',AZUL_OSC)
trow([
    ('Aplicacion de SOLID (backend)','3.0','5 principios con codigo ANTES/DESPUES y justificacion'),
    ('Refactorizacion documentada','1.5','Comparacion clara para cada principio con olor eliminado'),
    ('Tecnicas UI/UX (minimo 6)','2.0','8 tecnicas con justificacion Nielsen y capturas del sistema'),
    ('Funcionamiento del proyecto','1.5','Sistema funcional en Render.com'),
    ('Video de demostracion','1.5','(Por grabar por el estudiante)'),
    ('Documentacion / README','0.5','Este documento + README.md del repositorio'),
],['Criterio','Pts Maximos','Evidencia'])
callout('TOTAL MAXIMO: 10.0 puntos — El proyecto cumple con todos los criterios de la rubrica excepto el video de demostracion (pendiente de grabar).','EFF6FF')

# ============================================================
# FOOTER
# ============================================================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(20)
r=p.add_run('SIGEST v2.0 — Universidad Nacional Evangelica (UNEV)\nFrancisco R. Diaz — Julio 2026\nIngenieria de Software II (INF-259) — Tarea 4')
r.font.size=Pt(10); r.font.color.rgb=S500

doc.save(OUT_REAL)
print('Documento guardado en:', OUT_REAL)
